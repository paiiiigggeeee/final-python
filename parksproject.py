import requests
import random
import docx
from pprint import pprint

list_of_parks_url = 'https://mn-state-parks.herokuapp.com/api/list'
# not all information included in this api (name and park_id)
# view these as strings

# Make request to park list using the URL given as an argument
park_list_response = requests.get(list_of_parks_url).json()
pprint(park_list_response)

print("Choosing five state parks at random.")
park_list = random.sample(park_list_response, 5)  # Random sample of 5 parks is chosen from park list response
print(park_list)
print(len(park_list))

park_detail_api_url = 'https://mn-state-parks.herokuapp.com/api/'
document = docx.Document()  # Create word document
document.add_paragraph('Minnesota State Park Travel Guide', 'Title')  # Add paragraph with Title info

# Create for loop the retrieves list and dictionary data for all 5 parks from API
for park in park_list:
    park_id = park.get('park_id')  # Retrieve park_id using .get
    unique_park_api_url = park_detail_api_url + park_id  # create unique park api that adds detailed park api with each park id
    detailed_park_api = requests.get(unique_park_api_url).json()  # Make API request to new unique API
    # define the following variables in API and use .get to retrieve
    park_info = detailed_park_api.get('park_information')
    park_name = detailed_park_api.get('name')
    highlights = detailed_park_api.get('highlights')
    Address = detailed_park_api.get('address')
    Website = detailed_park_api.get('url')

    # Define park images variable and retrieve using .get
    park_images = detailed_park_api.get('park_images')

    document.add_heading(park_name)  # Add park name heading

    # Insert first photo from park images below park name heading
    title_photo = requests.get(park_images[0])  # request first indexed photo in park_images
    title_filename = park_id + 'park1.jpg'
    with open(title_filename, 'wb') as file:
        for chunk_of_binary_data in title_photo.iter_content():
            file.write(chunk_of_binary_data)
    document.add_picture(title_filename, width=docx.shared.Inches(3))  # Add photo to word document

    document.add_heading('Highlights')
    for highlight in highlights:  # for loop adds each individual highlight in highlights list
        document.add_paragraph(highlight, 'List Bullet')

    for category, text in park_info.items():  # for loop retrieves data in dictionary using .get
        geology = park_info.get('Geology')
        history = detailed_park_api.get('History')
        landscape = detailed_park_api.get('Landscape')
        wildlife = detailed_park_api.get('Wildlife')
        document.add_heading(category)  # heading is generated from category variable in dict
        document.add_paragraph(text)  # text is generated from text variable in dict

    # for loop gathers images from park_images variable
    for image in park_images:
        photo_response = requests.get(image)
        filename = park_id + 'park.jpg'  # create variable that adds park_id to the end of .jpg file
        with open(filename, 'wb') as file:
            for chunk_of_binary_data in photo_response.iter_content():
                file.write(chunk_of_binary_data)
        document.add_picture(filename, width=docx.shared.Inches(3))  # add pictures to word doc

    # Add final list data to word document
    document.add_paragraph('Contact Information', 'Heading 1')

    document.add_paragraph('Address', 'Heading 2')
    document.add_paragraph(Address)

    document.add_paragraph('Website', 'Heading 3')
    document.add_paragraph(Website)

document.save('detailedparks.docx')  # save document to the following file
